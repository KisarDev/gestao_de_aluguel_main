from datetime import datetime, timedelta
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from gestaoAluguel.models import Casa, Inquilino


def _gerar_contrato(usuario, id):
    """Faz uma query do inquilino selecionado pelo botão e substitui os valores do contrato pelos valores padrão"""
    inquilino = get_object_or_404(Inquilino, dono=usuario, id=id)
    data_inicio = datetime.now().strftime("%d/%m/%Y")
    data_fim = (datetime.now() + timedelta(days=180)
                ).strftime("%d/%m/%Y")  # 180 dias = 6 meses
    valor_aluguel = 900.00

    endereco = "XXXXXXXX"

    document = Document()

    titulo = document.add_paragraph("CONTRATO DE LOCAÇÃO")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.runs[0].font.size = Pt(16)
    titulo.runs[0].font.bold = True

    document.add_paragraph(f"""
Os signatários deste instrumento, de um lado Cesar Augusto Martins Silva, brasileiro, casado, residente nesta cidade de Extrema MG à rua Carolina Ferreira de Andrade, nº 30, Bairro Tenentes 4, Extrema MG, portador do RG 23.991.978 e CPF 142.374.006-89, e do outro lado {inquilino.nome}, brasileiro, {inquilino.profissao}, {inquilino.estado_civil}, portador do RG nº {inquilino.rg} , CPF Nº {inquilino.cpf}, residente e domiciliado a Rua Carolina Ferreira De Andrade, 30 Tenentes 4 - Extrema MG, tem justos e contratados o seguinte, que mutuamente convencionam, outorgam e aceitam, a saber:

1ª) O prazo de locação é de 6 meses (6 meses) a iniciar em {data_inicio} e a terminar em {data_fim}, data em que o locatário se obriga a restituir o imóvel desocupado ou de outra forma a renovar expressamente o novo contrato caso venha a permanecer no imóvel.

2ª) O valor do aluguel é de R$ {valor_aluguel:.2f} (novecentos reais) que O LOCATÁRIO se compromete a pagar pontualmente no dia 10 de cada mês a vencer, no estabelecimento residencial do LOCADOR ou de seu representante, mês a vencer.

3ª) Os consumos de água, luz, telefone e gás, assim como todos os encargos e tributos que indicam ou venham a incidir sobre o imóvel, conservação, seguro e outras decorrentes de Lei, assim como suas respectivas majorações, ficam a cargo do LOCATÁRIO.

4ª) O LOCATÁRIO, salvo as obras que importem na segurança do imóvel, obriga-se a todas as outras, devendo trazer o imóvel locado em boas condições de higiene e limpeza, e em perfeito estado de conservação e funcionamento, para assim restituí-lo quando findo ou rescindido este contrato sem direito a indenização por quaisquer benfeitorias ainda que necessárias, as quais ficarão desde logo incorporadas ao imóvel;

5ª) Obriga-se O LOCATÁRIO no curso de locação, a satisfazer todas as exigências dos Poderes Públicos a que der causa, não motivando elas a rescisão deste contrato;

6ª) Não é permitido a transferência deste contrato, nem a sub-locação, cessão ou empréstimo total ou parcial do imóvel, sem aviso prévio consentimento por escrito do LOCADOR, devendo no caso deste ser dado, agir oportunamente junto aos ocupantes, a fim de que o imóvel esteja desimpedido nos termos do presente contrato. Igualmente não é permitido fazer modificações, transformações no imóvel, sem autorização escrita do LOCADOR;

7ª) O LOCATÁRIO desde já faculta ao LOCADOR ou seu Representante, examinar ou vistoriar o imóvel locado quando entender conveniente;

8ª) No caso de desapropriação do imóvel locado, ficará o LOCADOR desobrigado por todas cláusulas deste contrato, ressalvado ao LOCATÁRIO, tão somente, a faculdade de haver no poder desapropriante a indenização a que, por ventura, tiver direito;

9ª) Nenhuma intimação do serviço sanitário será motivo para O LOCATÁRIO abandonar o imóvel ou pedir rescisão deste contrato, salvo procedendo vistoria judicial, que apure estar a construção ameaçando a ruir;

10ª) Para todas as questões oriundas deste contrato, será competente o Foro da situação do imóvel, com renúncia de qualquer outro, por mais especial que se apresente;

11ª) Tudo quanto for devido em razão do presente contrato e que não comportem o processo executivo, será cobrado em ação competente, ficando a cargo do devedor, em qualquer caso, os honorários advocatícios que o credor constituir para ressalva de seus direitos;

12ª) Quaisquer estragos ocasionados ao imóvel e suas instalações, bem como as despesas a que o proprietário for obrigado por eventuais modificações feitas no imóvel, ficam sob responsabilidade do LOCATÁRIO.

13ª) O locatário se compromete a pagar um mês de aluguel antecipado ao Locador.

14ª) Assinam o presente contrato em 02 vias, de igual teor, em presença das testemunhas abaixo, destinando-se uma via para cada uma das partes interessadas.

15ª) O volume do som, após às 22h00min deve ser baixo para não incomodar vizinhos, conforme determina a lei do Barulho Excessivo. Práticas de perturbação a vizinhança, provocando desordens, bem como quaisquer outras manifestações que venham a abalar a tranquilidade dos vizinhos e inclusive a reputação do local, devem ser evitadas e serão punidas, caso ocorram e venham a ser constatadas e objeto de reclamações e/ou boletins de ocorrência, responsabilizando-se o LOCATÁRIO e/ou infratores.

Extrema - MG, {datetime.now().strftime("%d de %B de %Y")}

Cesar Augusto Martins Silva

{inquilino.nome}

TESTEMUNHAS:

________________________   ______________________
    """)
    for paragrafo in document.paragraphs:
        for run in paragrafo.runs:
            run.font.name = "Arial"

    file_path = f"contrato_{inquilino.nome}.docx"
    document.save(file_path)

    with open(file_path, 'rb') as f:
        response = HttpResponse(f.read(
        ), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_path}'
        return response
